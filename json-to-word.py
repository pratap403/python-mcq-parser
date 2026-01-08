"""
Convert MCQ JSON to formatted Word document
"""
import json
import argparse
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


def create_mcq_document(json_path, output_path, title="MCQ Questions"):
    """Convert JSON MCQs to a formatted Word document"""
    
    # Load JSON
    with open(json_path, 'r', encoding='utf-8') as f:
        questions = json.load(f)
    
    print(f"📄 Loaded {len(questions)} questions from {json_path}")
    
    # Create Word document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Add title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add subtitle with stats
    subtitle = doc.add_paragraph(f"Total Questions: {len(questions)}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Blank line
    
    # Add each question
    for i, q in enumerate(questions, 1):
        add_question(doc, q, i)
    
    # Save document
    doc.save(output_path)
    print(f"✅ Saved to {output_path}")
    
    return len(questions)


def add_question(doc, q, index):
    """Add a single MCQ question to the document"""
    
    q_no = q.get('question_no', index)
    q_text = q.get('question', '')
    options = q.get('options', {})
    answer = q.get('answer', '')
    
    # Question number and text (bold number)
    q_para = doc.add_paragraph()
    q_run = q_para.add_run(f"Q{q_no}. ")
    q_run.bold = True
    q_para.add_run(q_text)
    
    # Options (each on new line, indented)
    for letter in ['A', 'B', 'C', 'D']:
        if letter in options:
            opt_text = options[letter]
            opt_para = doc.add_paragraph()
            opt_para.paragraph_format.left_indent = Inches(0.5)
            
            # Highlight correct answer
            if letter == answer:
                opt_run = opt_para.add_run(f"({letter}) {opt_text}")
                opt_run.bold = True
                # Add checkmark for correct answer
                opt_para.add_run(" ✓").bold = True
            else:
                opt_para.add_run(f"({letter}) {opt_text}")
    
    # Answer line
    if answer:
        ans_para = doc.add_paragraph()
        ans_para.paragraph_format.left_indent = Inches(0.5)
        ans_run = ans_para.add_run(f"Answer: {answer}")
        ans_run.bold = True
        ans_run.italic = True
    
    # Add spacing after question
    doc.add_paragraph()


def main():
    parser = argparse.ArgumentParser(description='Convert MCQ JSON to Word document')
    parser.add_argument('input', help='Input JSON file path')
    parser.add_argument('--output', '-o', help='Output Word file path (default: same name with .docx)')
    parser.add_argument('--title', '-t', default='MCQ Questions', help='Document title')
    
    args = parser.parse_args()
    
    # Determine output path
    if args.output:
        output_path = args.output
    else:
        output_path = args.input.rsplit('.', 1)[0] + '.docx'
    
    print(f"\n{'='*60}")
    print("JSON to Word Converter")
    print(f"{'='*60}\n")
    
    create_mcq_document(args.input, output_path, args.title)


if __name__ == "__main__":
    main()

